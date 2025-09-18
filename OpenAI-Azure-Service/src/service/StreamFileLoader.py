import logging
import os
import traceback
from io import BytesIO
from tempfile import NamedTemporaryFile
from typing import Any, List, Mapping, Optional

import fitz
import openpyxl
import pandas as pd
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain_community.document_loaders.pdf import BasePDFLoader
from langchain_core.documents.base import Blob
from opentelemetry import trace
from pymupdf import FileDataError

from src.service.documento_service import num_tokens_from_string
from src.util.upload_util import write_data_in_temporary_file

logger = logging.getLogger(__name__)
tracer = trace.get_tracer(__name__)


def _split_content_into_segments(content: str) -> List[str]:
    max_length = 100000
    segments = []
    current_segment = ""
    current_segment_length = 0

    for line in content.split("\n"):
        num_tokens = num_tokens_from_string(line)
        if current_segment_length + num_tokens < max_length:
            current_segment += line + "\n"
            current_segment_length += num_tokens
        else:
            segments.append(current_segment)
            current_segment = line + "\n"
            current_segment_length = num_tokens

    if current_segment:
        segments.append(current_segment)

    return segments


def _parse_pdf(data: BytesIO, file_name: str) -> str:
    with NamedTemporaryFile(suffix=file_name, delete=False) as temp_file:
        file_path = write_data_in_temporary_file(data, temp_file)

    try:
        with fitz.open(filename=file_path, filetype="pdf") as doc:
            content = []
            for page_number in range(doc.page_count):
                page = doc.load_page(page_number)
                if page_number == 0:
                    page_content = f"{file_name}: {page.get_text('text')}"
                else:
                    page_content = page.get_text("text")
                content.append(page_content)
            return "\n".join(content)
    except Exception as e:
        logger.error(f"Failed to open file '{file_path}' as type 'pdf': {e}")
        raise FileDataError(f"Failed to open file '{file_path}' as type 'pdf'.")
    finally:
        if file_path:
            os.remove(file_path)


def _parse_csv(data: BytesIO, file_name: str) -> str:
    df = pd.read_csv(data)
    content = df.to_csv(index=False)
    return file_name + ":" + content


def _parse_xlsx(data: BytesIO, file_name: str) -> str:
    wb = openpyxl.load_workbook(data)
    content = []
    for sheet in wb:
        for row in sheet.iter_rows(values_only=True):
            content.append("\t".join([str(cell) for cell in row if cell is not None]))
    content_str = "\n".join(content)
    return file_name + ":" + content_str


def _parse_docx(data: BytesIO, file_name: str) -> str:
    doc = DocxDocument(data)
    content = "\n".join([para.text for para in doc.paragraphs])
    return file_name + ":" + content


class StreamFileLoader(BasePDFLoader):
    """Load `PDF`, `DOCX`, `XLSX`, and `CSV` files."""

    @tracer.start_as_current_span("__init___BasePDFLoader")
    def __init__(
        self,
        text_kwargs: Optional[Mapping[str, Any]] = None,
        dedupe: bool = False,
        data_list: List[BytesIO] = None,
        file_names: List[str] = None,
    ) -> None:
        self.text_kwargs = text_kwargs or {}
        self.dedupe = dedupe
        self.data_list = data_list
        self.file_names = file_names

    @tracer.start_as_current_span("load_list")
    def load_list(self) -> List[Document]:
        logger.info(">> Executando o load_list dos documentos")
        try:
            blobs = []
            for data, file_name in zip(self.data_list, self.file_names):
                blobs.append(Blob.from_data(data.read(), path=file_name))

            all_content = []
            for blob in blobs:
                file_extension = blob.source.split(".")[-1].lower()
                if file_extension == "pdf":
                    content = _parse_pdf(BytesIO(blob.data), blob.source)
                elif file_extension == "docx":
                    content = _parse_docx(BytesIO(blob.data), blob.source)
                elif file_extension == "xlsx":
                    content = _parse_xlsx(BytesIO(blob.data), blob.source)
                elif file_extension == "csv":
                    content = _parse_csv(BytesIO(blob.data), blob.source)
                else:
                    continue

                all_content.extend(content)

            combined_content = "".join(all_content)
            final_segments = _split_content_into_segments(combined_content)

            documents = []
            for segment in final_segments:
                documents.append(
                    Document(page_content=segment, metadata={"source": "combined"})
                )

            return documents
        except Exception as e:
            logger.error(f"Erro ao processar o documento: {e}")
            traceback.print_exc()
            raise e
